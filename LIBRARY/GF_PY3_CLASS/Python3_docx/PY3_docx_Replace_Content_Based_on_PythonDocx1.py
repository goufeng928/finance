# PY3_docx_Replace_Content_Based_on_PythonDocx1.py
# Create by GF 2026-01-08 15:01

import docx  # python-docx 1.2.0

# ##################################################

class PY3_docx_Replace_Content_Based_on_PythonDocx1():

    def docx_File_Input(self, docx_Path:str):

        docx_File = docx.Document(docx = docx_Path)
        # ..........................................
        return docx_File

    def docx_Paragraphs_Replace_Text_Inplace(self, docx_Paragraphs:list, Old_Text:str, New_Text:str) -> int:

        i:int = 0
        while (i < len(docx_Paragraphs)):
            if (Old_Text in docx_Paragraphs[i].text):
                Saved_Paragraph_Style = docx_Paragraphs[i].style
                # ..................................
                j:int = 0
                while (j < len(docx_Paragraphs[i].runs)):
                    if (Old_Text in docx_Paragraphs[i].runs[j].text):
                        Saved_Run_font = docx_Paragraphs[i].runs[j].font
                        # ..........................
                        docx_Paragraphs[i].runs[j].text = docx_Paragraphs[i].runs[j].text.replace(Old_Text, New_Text)
                        # ..........................
                        docx_Paragraphs[i].runs[j].font.name      = Saved_Run_font.name
                        docx_Paragraphs[i].runs[j].font.size      = Saved_Run_font.size
                        docx_Paragraphs[i].runs[j].font.bold      = Saved_Run_font.bold
                        docx_Paragraphs[i].runs[j].font.italic    = Saved_Run_font.italic
                        docx_Paragraphs[i].runs[j].font.underline = Saved_Run_font.underline
                        docx_Paragraphs[i].runs[j].font.color.rgb = Saved_Run_font.color.rgb
                    # ..............................
                    j = j + 1
                # ..................................
                docx_Paragraphs[i].text = docx_Paragraphs[i].text.replace(Old_Text, New_Text)
                # ..................................
                docx_Paragraphs[i].style = Saved_Paragraph_Style
            # ......................................
            i = i + 1
        # ..........................................
        return 1

    def docx_File_Replace_Text_in_Paragraphs(self, docx_File, Old_Text:str, New_Text:str):

        docx_File_Copy = docx_File

        # docx.Document( ... ).paragraphs 输出示例 (List 类型):
        # >>> import docx
        # >>> docx_File = docx.Document(docx = docx_Path)
        # >>> print(docx_File.paragraphs)
        # [<docx.text.paragraph.Paragraph object at 0x0000029BC18992B0>,
        #  <docx.text.paragraph.Paragraph object at 0x0000029BC10B87D0>,
        #  ......
        #  <docx.text.paragraph.Paragraph object at 0x0000029BC131CCB0>]
        # ..........................................
        # docx.Document( ... ).paragraphs[i].style 输出示例:
        # >>> print(docx_File.paragraphs[1].style)
        # _ParagraphStyle('Body Text Indent') id: 2868009354080
        # ..........................................
        # docx.Document( ... ).paragraphs[i].runs 输出示例 (List 类型):
        # >>> print(docx_File.paragraphs[1].runs)
        # [<docx.text.run.Run object at 0x0000029BC2ADFF50>,
        #  <docx.text.run.Run object at 0x0000029BC2ADEF20>,
        #  ......
        #  <docx.text.run.Run object at 0x0000029BC2ADEEB0>]

        i:int = 0
        while (i < len(docx_File.paragraphs)):  # 遍历 docx 文档中每个 Paragraph
            if (Old_Text in docx_File.paragraphs[i].text):
                Saved_Paragraph_Style = docx_File.paragraphs[i].style  # 保存段落的样式
                # ..................................
                j:int = 0
                while (j < len(docx_File.paragraphs[i].runs)):  # 遍历段落中每个 Run (文本片段)
                    if (Old_Text in docx_File.paragraphs[i].runs[j].text):
                        Saved_Run_font = docx_File.paragraphs[i].runs[j].font  # 保存 Run 的字体样式
                        # ..........................
                        docx_File_Copy.paragraphs[i].runs[j].text = docx_File_Copy.paragraphs[i].runs[j].text.replace(Old_Text, New_Text)
                        # ..........................
                        docx_File_Copy.paragraphs[i].runs[j].font.name      = Saved_Run_font.name
                        docx_File_Copy.paragraphs[i].runs[j].font.size      = Saved_Run_font.size
                        docx_File_Copy.paragraphs[i].runs[j].font.bold      = Saved_Run_font.bold
                        docx_File_Copy.paragraphs[i].runs[j].font.italic    = Saved_Run_font.italic
                        docx_File_Copy.paragraphs[i].runs[j].font.underline = Saved_Run_font.underline
                        docx_File_Copy.paragraphs[i].runs[j].font.color.rgb = Saved_Run_font.color.rgb
                    # ..............................
                    j = j + 1
                # ..................................
                docx_File_Copy.paragraphs[i].text = docx_File_Copy.paragraphs[i].text.replace(Old_Text, New_Text)
                # ..................................
                docx_File_Copy.paragraphs[i].style = Saved_Paragraph_Style  # 保持原段落样式
            # ......................................
            i = i + 1
        # ..........................................
        return docx_File_Copy

    def docx_File_Replace_Text_in_Tables(self, docx_File, Old_Text:str, New_Text:str):

        docx_File_Copy = docx_File
        # ..........................................
        m:int = 0
        while (m < len(docx_File.tables)):
            # 遍历 docx 文档中的每个 Table
            # - docx_File.tables[m]
            #   docx_File.tables[m] 代表 docx 文档中第 m 个表格
            n:int = 0
            while (n < len(docx_File.tables[m].rows)):
                # # 遍历表格中每个 Row (表格行)
                # - docx_File.tables[m].rows[n]
                #   docx_File.tables[m].rows[n] 代表表格中第 n 行
                i:int = 0
                while (i < len(docx_File.tables[m].rows[n].cells)):
                    # 遍历表格行中每个 Cell (表格单元格)
                    # - docx_File.tables[m].rows[n].cells[i]
                    #   docx_File.tables[m].rows[n].cells[i] 代表表格行中第 i 个单元格
                    self.docx_Paragraphs_Replace_Text_Inplace(
                        # 提取表格单元格中每个 Paragraph (表格内嵌段落)
                        # - docx_File.tables[m].rows[n].cells[i].paragraphs[j]
                        #   docx_File.tables[m].rows[n].cells[i].paragraphs[j] 代表表格内嵌段落中第 j 个段落
                        docx_Paragraphs = docx_File_Copy.tables[m].rows[n].cells[i].paragraphs,
                        Old_Text        = Old_Text,
                        New_Text        = New_Text)
                    # ..............................
                    i = i + 1
                # ..................................
                n = n + 1
            # ......................................
            m = m + 1
        # ..........................................
        return docx_File_Copy

    def docx_File_Replace_Image_By_Image_Width_CM(self, docx_File, Image_Width_CM:float, New_Image_Path:str):

        docx_File_Copy = docx_File
        # ..........................................
        for Element in docx_File_Copy.element.body:
            # 处理 docx 段落
            if (Element.tag.endswith('p') == True):
                # 处理嵌套在 p 元素中的 drawing 元素 (图像以 "嵌入文本 (inline)" 方式插入)
                for drawing in Element.xpath(".//w:drawing"):  # 提取段落中的图像
                    blips:list = drawing.xpath(".//a:blip")
                    # ..............................
                    for blip in blips:
                        embed_id:str   = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        image_path:str = f"word/media/image{embed_id}.jpg"  # 图像在文档中的位置

                        # 获取图像尺寸信息
                        # 转换单位为EMU（English Metric Units）
                        # 1英寸 = 914400 EMU, 96 DPI下1像素 ≈ 9525 EMU
                        # 1厘米 = 360000 EMU
                        # 1毫米 = 36000 EMU
                        # 1磅（point）= 12700 EMU（1英寸=72磅）
                        # 但更准确的是保持原有的EMU值
                        extents = drawing.xpath(".//wp:extent")
                        if extents:
                            cx = extents[0].get("cx")  # 宽度 (EMU 单位)
                            cy = extents[0].get("cy")  # 高度 (EMU 单位)
                        else:
                            cx = None
                            cy = None

                        if (round(int(cx) / 360000, 2) == Image_Width_CM):

                            # 查找对应的 Relationship
                            if (embed_id in docx_File_Copy.part.related_parts):
                                image_part       = docx_File_Copy.part.related_parts[embed_id]
                                image_data:bytes = image_part._blob
                                image_type:str   = image_part.content_type
                            
                            with open(New_Image_Path, 'rb') as f:
                                new_image_data = f.read()
                            image_part._blob = new_image_data
            
            # 处理 docx 文档主体 (body) 首层的绘图元素
            # 当图像以 "浮动" 方式插入, 而不是 "嵌入文本" 时, 可能会作为独立对象
            if (Element.tag.endswith("drawing") == True):    
                blips:list = drawing.xpath(".//a:blip")
                # ..................................
                for blip in blips:
                    embed_id:str   = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    image_path:str = f"word/media/image{embed_id}.jpg"  # 图像在文档中的位置

                    extents = drawing.xpath(".//wp:extent")
                    if extents:
                        cx = extents[0].get("cx")  # 宽度 (EMU 单位)
                        cy = extents[0].get("cy")  # 高度 (EMU 单位)
                    else:
                        cx = None
                        cy = None
                    
                    if (round(int(cx) / 360000, 2) == Image_Width_CM):
                    
                        # 查找对应的 Relationship
                        if (embed_id in docx_File_Copy.part.related_parts):
                            image_part       = docx_File_Copy.part.related_parts[embed_id]
                            image_data:bytes = image_part._blob
                            image_type:str   = image_part.content_type
                        
                        with open(New_Image_Path, 'rb') as f:
                            new_image_data = f.read()
                        image_part._blob = new_image_data
        # ..........................................
        return docx_File_Copy

# EOF Signed by GF.

# GF_PY3_CLASS/Python3_docx/PY3_docx_to_PDF_Based_on_ReportLab4.py
# Create by GF 2026-01-06 22:06

import io
import os
# ..................................................
import docx
import PIL
# ..................................................
from reportlab.lib             import colors
from reportlab.lib.pagesizes   import letter, A4
from reportlab.lib.styles      import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units       import inch
from reportlab.platypus        import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.pdfbase         import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ##################################################

class PY3_ReportLab4_docx(object):

    def PythonDocx1_Extract_Contents_as_List(self, docx_Path:str) -> list:

        docx_File = docx.Document(docx_Path)

        # docx.Document( ... ) 输出示例:
        # >>> docx_File = docx.Document(docx_Path)
        # >>> print(docx_Path)
        # <docx.document.Document object at 0x000001D7BE15C6E0>

        docx_Contents_List:list = []

        # XPath 表达式分解
        # "." 当前节点: 表示从当前元素开始, 如果不加 . 例如 "//w:r", 表示从文档根节点开始搜索
        # "//" 递归下降: 表示在当前节点及其所有后代节点中搜索, 不管元素在多少层嵌套下都会被找到
        # "w:r" 带命名空间的元素: "w:" 是命名空间前缀, "r" 是元素名 "w:r" 表示 "w" 命名空间下的 "r" 元素
        # ..........................................
        # 常见的 DOCX XML 元素
        # XPath 表达式  含义        说明
        # -------------------------------------------------------
        # .//w:p        段落        文档的基本文本单元
        # .//w:r        文本运行    应用样式的最小文本单位
        # .//w:t        文本        实际的文本内容
        # .//w:rPr      运行属性    文本样式 (字体, 大小, 颜色等)
        # .//w:pPr      段落属性    段落样式 (对齐, 缩进, 间距等)
        # .//w:tbl      表格        表格元素
        # .//w:tr       表格行      表格中的一行
        # .//w:tc       表格单元格  表格单元格
        # .//w:drawing  绘图        包含图片, 形状等
        # ..........................................
        # docx.Document( ... ).element.body 中 element 的 xpath(".//w:r") 输出示例:
        # >>> for Element in docx_File.element.body:
        # ...     print(Element.xpath('.//w:r'))
        # [<CT_R '<w:r>' at 0x2008019bf50>]
        # ......
        # []
        # [<CT_R '<w:r>' at 0x2008019b8f0>]

        for Element in docx_File.element.body:

            # docx.Document( ... ).element.body 输出示例:
            # >>> docx_File = docx.Document(docx_Path)
            # >>> print(docx_File.element.body)
            # <CT_Body '<w:body>' at 0x1d7bf3128d0>

            # docx.Document( ... ).element.body 中 element 的 tag 输出示例:
            # >>> for Element in docx_File.element.body:
            # ...     print(Element.tag)
            # {http://schemas.openxmlformats.org/wordprocessingml/2006/main}p
            # {http://schemas.openxmlformats.org/wordprocessingml/2006/main}p
            # ......
            # {http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl
            # {http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr

            # 处理 docx 段落
            if (Element.tag.endswith('p') == True):

                for run in Element.xpath(".//w:r"):
                    Text_Elements:list = run.xpath(".//w:t")  # 提取段落中的文本
                    Text:str           = str('')
                    # ..............................
                    if (Text_Elements != []):
                        for t in Text_Elements:
                            Text = "%s%s" % (Text, t.text)
                    # ..............................
                    docx_Contents_List.append({"mime": "text/plain", "text": Text})

                # 处理嵌套在 p 元素中的 drawing 元素 (图像以 "嵌入文本" 方式插入)
                for drawing in Element.xpath(".//w:drawing"):  # 提取段落中的图像
                    blips:list = drawing.xpath(".//a:blip")
                    # ..............................
                    for blip in blips:
                        embed_id:str   = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        image_path:str = f"word/media/image{embed_id}.jpg"  # 图像在文档中的位置
                        # ..........................
                        if (embed_id in docx_File.part.related_parts):
                            image_part       = docx_File.part.related_parts[embed_id]
                            image_data:bytes = image_part._blob
                            image_type:str   = image_part.content_type
                            # ......................
                            docx_Contents_List.append({"mime": image_type, "blob": image_data})

            # 处理 docx 文档主体 (body) 首层的绘图元素
            # 当图像以 "浮动" 方式插入, 而不是 "嵌入文本" 时, 可能会作为独立对象
            if (Element.tag.endswith("drawing") == True):

                blips:list = drawing.xpath(".//a:blip")
                # ..................................
                for blip in blips:
                    embed_id:str   = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    image_path:str = f"word/media/image{embed_id}.jpg"  # 图像在文档中的位置
                    # ..............................
                    if (embed_id in docx_File.part.related_parts):
                        image_part       = docx_File.part.related_parts[embed_id]
                        image_data:bytes = image_part._blob
                        image_type:str   = image_part.content_type
                        # ..........................
                        docx_Contents_List.append({"mime": image_type, "blob": image_data})

            # 处理 docx 表格
            if (Element.tag.endswith("tbl") == True):

                table_data:list     = []
                table_rows_xml:list = Element.xpath(".//w:tr")  # 获取当前表格所有行
                # ..................................
                for row_xml in table_rows_xml:
                    row_cells_xml:list = row_xml.xpath(".//w:tc")  # 获取当前行所有单元格
                    table_rows:list    = []
                    # ..............................
                    for cell_xml in row_cells_xml:
                        cell_text = ""
                        text_runs = cell_xml.xpath(".//w:t")  # 提取单元格中的文本
                        # ..........................
                        for t in text_runs:
                            cell_text = "%s%s" % (cell_text, t.text)
                        # ..........................
                        table_rows.append(cell_text)
                    # ..............................
                    table_data.append(table_rows)
                # ..................................
                docx_Contents_List.append(
                    {"mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.table",
                     "data": table_data}
                )
        # ..........................................
        return docx_Contents_List

    def Pillow8_Image_Keep_Ratio_Adjust_Width(self, Image_File:str, Width:float) -> float:

        Img = PIL.Image.open(Image_File)

        # Pillow 中 Image.open( ... ).size 的输出示例:
        # >>> from PIL import Image
        # >>> Img = Image.open("example.png")
        # >>> print(Img.size)
        # >>> Img.close()
        # (480, 270)

        Orig_Width, Orig_Height = Img.size
        # ..........................................
        Width_to_Height_Ratio:float = Orig_Width / Orig_Height
        # ..........................................
        Img.close()
        # ..........................................
        return Width / Width_to_Height_Ratio

    def ReportLab4_Register_Chinese_Font(self) -> str:

        Font_Name:str       = "Helvetica"  # ReportLab 4 默认字体
        Font_Path_List:list = ["C:/Windows/Fonts/simhei.ttf",   # Windows 系统 "黑体"
                               "C:/Windows/Fonts/simsun.ttc",   # Windows 系统 "宋体"
                               "C:/Windows/Fonts/simkai.ttf",   # Windows 系统 "楷体"
                               "/System/Library/Fonts/STHeiti Light.ttc",  # Mac 系统 "黑体"
                               "/System/Library/Fonts/STSong.ttf",         # Mac 系统 "宋体"
                               "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf"]  # Linux 系统

        i:int = 0
        while (i < len(Font_Path_List)):
            Font_Path:str = Font_Path_List[i]
            # ......................................
            if (os.path.exists(Font_Path) == True):
                Registering:str = os.path.basename(Font_Path).split('.')[0]
                # ..................................
                try:
                    pdfmetrics.registerFont(TTFont(Registering, Font_Path))
                    # ..............................
                    print(f"[DEBUG] reportlab.pdfbase.pdfmetrics.registerFont( ... ): {Registering}")
                    # ..............................
                    Font_Name = Registering
                    # ..............................
                    i = len(Font_Path_List)  # 立即达到循环停止条件
                except Exception as e:
                    print(f"[DEBUG] reportlab.pdfbase.pdfmetrics.registerFont( ... ): {e}")
            # ......................................
            i = i + 1
        # ..........................................
        return Font_Name

    def ReportLab4_docx_File_to_PDF_File(self, docx_Path:str, PDF_Path:str) -> int:

        docx_Contents_List:list = self.PythonDocx1_Extract_Contents_as_List(docx_Path = docx_Path)

        # reportlab.platypus 的 SimpleDocTemplate( ... ) 输出示例:
        # >>> from reportlab.platypus import SimpleDocTemplate
        # >>> PDF_Path = "example.pdf"
        # >>> PDF_File = PDF_File  = SimpleDocTemplate(PDF_Path, pagesize = letter)
        # >>> print(PDF_File)
        # <reportlab.platypus.doctemplate.SimpleDocTemplate object at 0x000001D7BE165E50>

        PDF_File = SimpleDocTemplate(filename     = PDF_Path,
                                     pagesize     = A4,  # 使用 A4 纸型
                                     rightMargin  = 72,
                                     leftMargin   = 72,
                                     topMargin    = 72,
                                     bottomMargin = 72)

        PDF_Story:list    = []
        PDF_Font_Name:str = self.ReportLab4_Register_Chinese_Font()
        PDF_Styles        = getSampleStyleSheet()

        for Line in docx_Contents_List:

            if (Line["mime"] == "text/plain"):
                PDF_Styles_Zh = ParagraphStyle(name           = "NormalChinese",
                                               parent         = PDF_Styles["Normal"],
                                               fontName       = "simhei.ttf",
                                               fontSize       = 10,
                                               leading        = 14,     # 行距
                                               spaceAfter     = 6,      # 段后间距
                                               alignment      = 4,
                                               wordWrap       = "CJK",  # 允许中文换行, 但西文可能需要额外处理
                                               splitLongWords = True,   # 允许长单词分割
                                               allowOrphans   = 0,      # 不允许孤行 (Orphans)
                                               allowWidows    = 0)      # 不允许寡行 (Widows)
                # ..................................
                PDF_Story.append(Paragraph(Line["text"], PDF_Styles_Zh))

            if (Line["mime"] == "image/jpeg"):
                Img_File         = io.BytesIO(Line["blob"])
                Img_Width:float  = 400
                Img_Height:float = self.Pillow8_Image_Keep_Ratio_Adjust_Width(Img_File, Img_Width)
                Img_File         = io.BytesIO(Line["blob"])
                # ..................................
                PDF_Story.append(Image(filename = Img_File, width = Img_Width, height = Img_Height))

            if (Line["mime"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.table"):
                PDF_Table = Table(Line["data"], colWidths = [100, 60, 100])
                # ..................................
                PDF_Table_Styles = TableStyle([('BACKGROUND',    (0, 0), (-1,  0), colors.grey),        # 表头背景
                                               ('TEXTCOLOR',     (0, 0), (-1,  0), colors.whitesmoke),  # 表头文字颜色
                                               ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),           # 居中对齐
                                               ('FONTNAME',      (0, 0), (-1,  0), 'Helvetica-Bold'),   # 表头字体
                                               ('FONTSIZE',      (0, 0), (-1,  0), 12),                 # 表头字号
                                               ('BOTTOMPADDING', (0, 0), (-1,  0), 12),                 # 表头底部间距
                                               ('BACKGROUND',    (0, 1), (-1, -1), colors.white),       # 数据行背景
                                               ('GRID',          (0, 0), (-1, -1), 1, colors.black)])   # 网格线
                # ..................................
                PDF_Table.setStyle(PDF_Table_Styles)
                # ..................................
                PDF_Story.append(PDF_Table)  # 添加到 PDF 文档流

        PDF_File.build(PDF_Story)
        # ..........................................
        return 1

# EOF Signed by GF
